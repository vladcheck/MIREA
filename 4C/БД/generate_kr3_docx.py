import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def clear_paragraph(p):
    p._element.getparent().remove(p._element)

def add_toc(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')  # Create a new w:fldChar element
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)

doc = Document('kr2/КР2_БД_Отчет_Валекжанин_В_С.docx')

if len(doc.paragraphs) > 3:
    doc.paragraphs[3].text = doc.paragraphs[3].text.replace("Контрольная работа № 2", "Контрольная работа № 3")
if len(doc.paragraphs) > 7:
    doc.paragraphs[7].text = "Настройка и поддержка транзакций в СУБД PostgreSQL"
if len(doc.paragraphs) > 8:
    doc.paragraphs[8].text = ""

# Delete from paragraph 23
for p in list(doc.paragraphs)[23:]:
    clear_paragraph(p)

for tbl in doc.tables:
    tbl._element.getparent().remove(tbl._element)

doc.add_page_break()

# Add TOC
toc_heading = doc.add_paragraph("Оглавление", style="Heading 1")
toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_p = doc.add_paragraph()
add_toc(toc_p)

doc.add_page_break()

def add_heading(text, style_name):
    try:
        p = doc.add_paragraph(text, style=style_name)
    except:
        p = doc.add_paragraph(text)
    if style_name == "Heading 1":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_normal(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    try:
        p.style = doc.styles['Normal']
    except:
        pass
    return p

def add_screenshot_placeholder(num, comment):
    p = doc.add_paragraph(f"[Скриншот {num} - {comment}]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    doc.add_paragraph("")

parts = [
    {
        "title": "Часть 1. Вставка, обновление и удаление строки",
        "goal": "Цель: Изучение работы с DML-запросами в транзакциях и исследование версионирования строк (xmin, xmax).",
        "file": "kr3/part1.sql",
        "shots": [
            "Вывод таблицы с xmin",
            "Вывод SELECT в обоих терминалах при незафиксированном UPDATE",
            "Вывод SELECT при незафиксированном DELETE",
            "Вывод после COMMIT",
            "Вывод", "Вывод", "Вывод"
        ]
    },
    {
        "title": "Часть 2. Видимость версии строки на различных уровнях изоляции",
        "goal": "Цель: Исследование уровней изоляции READ COMMITTED и REPEATABLE READ и их влияния на видимость изменений между транзакциями.",
        "file": "kr3/part2.sql",
        "shots": [
            "SHOW transaction_isolation",
            "Скриншот аномалии неповторяющегося чтения",
            "Скриншот уровня REPEATABLE READ без аномалии"
        ]
    },
    {
        "title": "Часть 3. Состояние транзакции по CLOG",
        "goal": "Цель: Изучить статус транзакций (in progress, committed, aborted) и работу механизма фиксации транзакций.",
        "file": "kr3/part3.sql",
        "shots": ["Статус in progress", "Статус committed", "Статус aborted"]
    },
    {
        "title": "Часть 4. Блокировки таблицы",
        "goal": "Цель: Рассмотреть блокировки таблиц (RowExclusiveLock и ShareLock) и возникновение ситуаций ожидания.",
        "file": "kr3/part4.sql",
        "shots": ["Вывод блокировок RowExclusiveLock", "Терминал 2 в ожидании ShareLock", "Успешное создание индекса"]
    }
]

scr_idx = 1
for idx, part in enumerate(parts):
    # Using Heading 2 per KR2 style
    add_heading(part["title"], "Heading 2")
    add_normal(part["goal"])
    
    try:
        with open(part["file"], "r", encoding="utf-8") as f:
            code = f.read()
        
        # Heading 3 per KR2 style (like 'Код функции' or 'Тесты')
        add_heading("Код", "Heading 3")
        add_normal(code)
    except FileNotFoundError:
        add_normal("Код не найден: " + part["file"])
        
    add_heading("Скриншоты выполнения", "Heading 3")
    for s_comment in part["shots"]:
        add_screenshot_placeholder(scr_idx, s_comment)
        scr_idx += 1
        
    if idx < len(parts) - 1:
        doc.add_page_break()

doc.save('kr3/КР3_БД_Отчет_Валекжанин_В_С.docx')
